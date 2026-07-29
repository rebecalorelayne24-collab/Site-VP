import io
import json
import os
import re
from datetime import datetime, timedelta
from zoneinfo import ZoneInfo

import docx
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Inches, Pt, RGBColor
import google.generativeai as genai
import pandas as pd
from pypdf import PdfReader
import streamlit as st

FUSO_BR = ZoneInfo("America/Sao_Paulo")

# Tabela oficial de preços base da Farmácia Jr. atualizada
PRECOS_AUTORAIS = {
    "Rotulagem Nutricional": 130.00,
    "Tabela Nutricional": 90.00,
    "Revisão Bibliográfica": 1100.00,
    "Formulação Cosmética": 250.00,
}


def obter_agora_br():
    """Retorna a data e hora atuais no fuso horário de Brasília."""
    return datetime.now(FUSO_BR)


def calcular_data_final_uteis(data_inicial, dias_uteis):
    """Calcula a data final pulando sábados e domingos automaticamente"""
    data_atual = data_inicial
    dias_contados = 0
    while dias_contados < dias_uteis:
        data_atual += timedelta(days=1)
        if data_atual.weekday() < 5:  # Segunda a sexta-feira
            dias_contados += 1
    return data_atual


@st.cache_data(show_spinner=False)
def _processar_pdf_ia_cached(bytes_pdf_content, api_key):
    """Processa o PDF usando o Gemini em cache para economizar chamadas de API."""
    genai.configure(api_key=api_key)

    reader = PdfReader(io.BytesIO(bytes_pdf_content))
    texto_completo = ""
    for page in reader.pages:
        texto_extraido = page.extract_text()
        if texto_extraido:
            texto_completo += texto_extraido + "\n"

    if not texto_completo.strip():
        return None, "Não foi possível extrair texto legível do PDF enviado."

    prompt = """
    Você é a inteligência do sistema financeiro da Farmácia Jr. (UFMG). 
    Analise o texto extraído de um PDF de orçamento de laboratório parceiro e encontre o VALOR TOTAL BRUTO do serviço.
    Retorne ESTRITAMENTE um JSON no seguinte formato, sem formatação markdown adicional ou blocos de código:
    {"valor_total": 0.00}
    """

    model = genai.GenerativeModel("gemini-2.5-flash")
    response = model.generate_content(
        f"{prompt}\n\nTexto do PDF:\n{texto_completo}",
        generation_config=genai.types.GenerationConfig(temperature=0.0),
    )

    texto_limpo = (
        response.text.strip().replace("```json", "").replace("```", "")
    )
    dados_ia = json.loads(texto_limpo)
    return float(dados_ia.get("valor_total", 0.0)), None


def extrair_valor_pdf_com_ia(arquivo_pdf):
    """Lê o PDF do orçamento do laboratório e extrai o valor utilizando o cache para economizar chamadas."""
    try:
        api_key = st.secrets.get("GEMINI_API_KEY") or os.environ.get(
            "GEMINI_API_KEY"
        )
        if api_key:
            api_key = str(api_key).strip().strip('"').strip("'")

        if not api_key:
            st.error(
                "API Key do Gemini não configurada. Adicione 'GEMINI_API_KEY'"
                " nos secrets."
            )
            return None

        arquivo_pdf.seek(0)
        bytes_pdf = arquivo_pdf.read()

        val, erro_msg = _processar_pdf_ia_cached(bytes_pdf, api_key)
        if erro_msg:
            st.warning(f"⚠️ {erro_msg}")
            return None

        return val

    except Exception as e:
        erro_str = str(e)
        if "429" in erro_str or "RESOURCE_EXHAUSTED" in erro_str:
            st.warning(
                "⏳ Cota gratuita da API atingida temporariamente. Digite o valor do laboratório no campo abaixo."
            )
        else:
            st.error(f"Erro ao processar o PDF com a IA: {e}")
        return None


def obter_texto_parcelamento(servico, valor_total):
    """Aplica as regras da planilha e calcula o valor de cada parcela"""
    if servico == "Revisão Bibliográfica":
        return "Consulte o diretor"

    num_parcelas = 1
    texto_especial = None

    if servico in ["Rotulagem Nutricional", "Rotulagem de cosméticos"]:
        if valor_total < 200:
            num_parcelas = 1
        elif valor_total < 600:
            num_parcelas = 2
        elif 600 <= valor_total < 800:
            num_parcelas = 3
        elif 800 <= valor_total <= 1000:
            num_parcelas = 4
        else:
            texto_especial = "Consulte o diretor"

    elif servico == "Tabela Nutricional":
        if valor_total < 200:
            num_parcelas = 1
        elif valor_total < 600:
            num_parcelas = 2
        elif 600 <= valor_total < 800:
            num_parcelas = 3
        elif 800 <= valor_total <= 1000:
            num_parcelas = 4
        else:
            texto_especial = (
                "Parcelamento Especial (Olhar com Presidência/Diretoria)"
            )

    if texto_especial:
        return texto_especial

    if num_parcelas == 1:
        v_str = f"R$ {valor_total:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
        return f"à vista ({v_str})"
    else:
        valor_da_parcela = valor_total / num_parcelas
        v_parc_str = f"R$ {valor_da_parcela:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
        v_tot_str = f"R$ {valor_total:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
        return f"à vista ({v_tot_str}) ou {num_parcelas}X de {v_parc_str}"


def definir_borda_celula(cell, **kwargs):
    """Aplica bordas finas customizadas nas células da tabela do Word."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>\n'
        f'<w:top w:val="{kwargs.get("top", "single")}" w:sz="{kwargs.get("sz", "4")}" w:space="0" w:color="{kwargs.get("color", "D3D3D3")}"/>\n'
        f'<w:left w:val="{kwargs.get("left", "single")}" w:sz="{kwargs.get("sz", "4")}" w:space="0" w:color="{kwargs.get("color", "D3D3D3")}"/>\n'
        f'<w:bottom w:val="{kwargs.get("bottom", "single")}" w:sz="{kwargs.get("sz", "4")}" w:space="0" w:color="{kwargs.get("color", "D3D3D3")}"/>\n'
        f'<w:right w:val="{kwargs.get("right", "single")}" w:sz="{kwargs.get("sz", "4")}" w:space="0" w:color="{kwargs.get("color", "D3D3D3")}"/>\n'
        f"</w:tcBorders>"
    )
    tcPr.append(tcBorders)


def formatar_tabela_word(tabela, bg_cabecalho="003366"):
    """Aplica o fundo Azul Escuro (003366) no cabeçalho e texto branco nas tabelas."""
    tabela.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row in enumerate(tabela.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if i == 0:
                # Fundo Azul para o cabeçalho
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg_cabecalho}"/>')
                cell._tc.get_or_add_tcPr().append(shading)
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for r in p.runs:
                        r.font.bold = True
                        r.font.color.rgb = RGBColor(255, 255, 255)  # Texto Branco
                        r.font.size = Pt(10)
            else:
                definir_borda_celula(cell)
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(10)


def gerar_docx_proposta(dados):
    """Gera o documento Word (.docx) com título PRETO e detalhes em AZUL."""
    doc = docx.Document()

    # Margens padrão (2.54 cm / 1 polegada)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(11)

    agora = obter_agora_br()
    data_validade = calcular_data_final_uteis(agora, 15).strftime("%d/%m/%Y")

    # 1. Validade (Topo Alinhado à Direita, tom cinza discreto)
    p_val = doc.add_paragraph()
    p_val.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_val = p_val.add_run(f"Validade precificação:\n{data_validade}")
    run_val.font.size = Pt(9.5)
    run_val.font.color.rgb = RGBColor(100, 100, 100)

    # 2. Título do Cliente (Centralizado, Negrito e em PRETO)
    p_tit = doc.add_paragraph()
    p_tit.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_tit = p_tit.add_run(f"\nPrecificação – {dados['nome_lead']}")
    run_tit.bold = True
    run_tit.font.size = Pt(15)
    run_tit.font.color.rgb = RGBColor(0, 0, 0)  # TÍTULO EM PRETO

    # 3. Subtítulo (Nome do Serviço em Azul)
    p_sub = doc.add_paragraph()
    run_sub = p_sub.add_run(f"{dados['nome_servico']}")
    run_sub.bold = True
    run_sub.font.size = Pt(12)
    run_sub.font.color.rgb = RGBColor(0, 51, 102)  # SUBTÍTULO EM AZUL

    if dados["tipo_servico"] == "Serviço Autoral (Farmácia Jr.)":
        # ---------------------------------------------------------------------
        # 1ª OPÇÃO - PRECIFICAÇÃO CHEIA
        # ---------------------------------------------------------------------
        p_op1 = doc.add_paragraph()
        run_op1 = p_op1.add_run("1° opção – Precificação cheia")
        run_op1.bold = True
        run_op1.font.color.rgb = RGBColor(0, 51, 102)  # Título de opção em Azul

        tbl1 = doc.add_table(rows=2, cols=4)
        tbl1.autofit = False

        hdr1 = tbl1.rows[0].cells
        hdr1[0].text, hdr1[1].text, hdr1[2].text, hdr1[3].text = (
            "Serviços",
            "Valor unitário",
            "Quantidade proposta",
            "Total",
        )

        row1 = tbl1.rows[1].cells
        row1[0].text = dados["nome_servico"]
        row1[1].text = f"R$ {dados['valor_base']:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
        row1[2].text = str(dados["quantidade"])
        row1[3].text = f"R$ {dados['total_cheio']:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")

        formatar_tabela_word(tbl1, bg_cabecalho="003366")  # Cabeçalho da Tabela em Azul

        doc.add_paragraph(f"\nPrazo de execução: {dados['prazo']} dias úteis")
        doc.add_paragraph(f"Formas de pagamento: {dados['parcelas_cheio']}")

        doc.add_paragraph()  # Espaçador

        # ---------------------------------------------------------------------
        # 2ª OPÇÃO - DESCONTO DE ACORDO
        # ---------------------------------------------------------------------
        p_op2 = doc.add_paragraph()
        run_op2 = p_op2.add_run("2° opção – desconto de acordo")
        run_op2.bold = True
        run_op2.font.color.rgb = RGBColor(0, 51, 102)  # Título de opção em Azul

        # Marcadores de desconto em tópicos
        if dados.get("motivos_lista") and isinstance(dados["motivos_lista"], list):
            for motivo in dados["motivos_lista"]:
                p_item = doc.add_paragraph(style="List Bullet")
                p_item.add_run(motivo)
        elif dados.get("motivo_desconto"):
            for motivo in dados["motivo_desconto"].split(", "):
                if motivo.strip():
                    p_item = doc.add_paragraph(style="List Bullet")
                    p_item.add_run(motivo.strip())

        tbl2 = doc.add_table(rows=2, cols=4)
        tbl2.autofit = False

        hdr2 = tbl2.rows[0].cells
        hdr2[0].text, hdr2[1].text, hdr2[2].text, hdr2[3].text = (
            "Serviços",
            "Valor unitário",
            "Quantidade proposta",
            "Total",
        )

        row2 = tbl2.rows[1].cells
        row2[0].text = dados["nome_servico"]
        row2[1].text = f"R$ {dados['unitario_desconto']:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
        row2[2].text = str(dados["quantidade"])
        row2[3].text = f"R$ {dados['total_desconto']:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")

        formatar_tabela_word(tbl2, bg_cabecalho="003366")  # Cabeçalho da Tabela em Azul

        doc.add_paragraph(f"\nPrazo de execução: {dados['prazo']} dias úteis")
        doc.add_paragraph(f"Formas de pagamento: {dados['parcelas_desconto']}")

    else:
        # ---------------------------------------------------------------------
        # TERCEIRIZADO (LABORATÓRIO)
        # ---------------------------------------------------------------------
        p_op1 = doc.add_paragraph()
        run_op1 = p_op1.add_run("1° opção – Precificação cheia")
        run_op1.bold = True
        run_op1.font.color.rgb = RGBColor(0, 51, 102)

        tbl3 = doc.add_table(rows=2, cols=4)
        tbl3.autofit = False

        hdr3 = tbl3.rows[0].cells
        hdr3[0].text, hdr3[1].text, hdr3[2].text, hdr3[3].text = (
            "Serviços",
            "Valor unitário",
            "Quantidade proposta",
            "Total",
        )

        row3 = tbl3.rows[1].cells
        row3[0].text = dados["nome_servico"]
        row3[1].text = f"R$ {dados['total_terceirizado']:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
        row3[2].text = "1"
        row3[3].text = f"R$ {dados['total_terceirizado']:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")

        formatar_tabela_word(tbl3, bg_cabecalho="003366")

        doc.add_paragraph(
            f"\nPrazo de execução: {dados['prazo_terceirizado']} dias (Incluso prazo de segurança da EJ. Previsão: {dados['data_entrega_terc']})"
        )
        doc.add_paragraph(f"Formas de pagamento: {dados['parcelas_terceirizado']}")

        doc.add_paragraph()

        p_param = doc.add_paragraph()
        p_param.add_run("Parâmetros analisados:\n").bold = True
        p_param.add_run(dados["parametros"])

        p_met = doc.add_paragraph()
        p_met.add_run("\nMetodologia:\n").bold = True
        p_met.add_run(dados["metodologia"])

        doc.add_paragraph(f"\nValor da coleta: {dados['txt_coleta']}")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def renderizar_aba_precificacao():
    st.title("🦩 Calculadora de Precificação Comercial Inteligente")

    st.markdown(
        """
        <div style="background-color: #F0F8FF; padding: 15px; border-left: 5px solid #003366; border-radius: 4px; margin-bottom: 20px;">
            <span style="color: #003366; font-weight: bold;">⚠️ DIRETRIZ INTERNA DA VICE-PRESIDÊNCIA:</span><br>
            O prazo de entrega de uma precificação de um <b>serviço autoral</b> para negócios é de no máximo <b>2 dias úteis</b>.<br>
            Já em caso de <b>serviço laboratorial (terceirizado)</b> não podemos exceder o prazo máximo de <b>5 dias úteis</b> para envio ao lead.
        </div>
    """,
        unsafe_allow_html=True,
    )

    nome_lead = st.text_input(
        "Nome do Lead / Empresa:", placeholder="Ex: Mika Doces Artesanais"
    )
    tipo_servico = st.selectbox(
        "Modalidade do Serviço:",
        [
            "Serviço Autoral (Farmácia Jr.)",
            "Serviço Terceirizado (Laboratório)",
        ],
    )

    dados_calculados = {}
    agora = obter_agora_br()

    if tipo_servico == "Serviço Autoral (Farmácia Jr.)":
        col1, col2 = st.columns(2)
        with col1:
            nome_servico = st.selectbox(
                "Serviço Autoral:", list(PRECOS_AUTORAIS.keys())
            )
            valor_base = PRECOS_AUTORAIS[nome_servico]
            st.info(
                f"💰 Valor de tabela fixado pelo setor: **R$ {valor_base:,.2f}**"
            )
            quantidade = st.number_input(
                "Quantidade (Nº de Rótulos / Tabelas / Produtos):",
                min_value=1,
                value=1,
            )
        with col2:
            if nome_servico == "Rotulagem Nutricional":
                prazo = quantidade * 10 if quantidade <= 1 else quantidade * 6
            elif nome_servico == "Tabela Nutricional":
                prazo = quantidade * 7 if quantidade <= 1 else quantidade * 4
            elif nome_servico == "Revisão Bibliográfica":
                prazo = quantidade * 60
            else:
                prazo = quantidade * 5

            st.info(
                f"📅 Prazo de execução calculated: **{prazo} dias úteis**"
            )

            total_cheio = valor_base * quantidade
            parcelas_cheio = obter_texto_parcelamento(
                nome_servico, total_cheio
            )
            st.info(f"💳 Condição de Pagamento (Cheio): **{parcelas_cheio}**")

        data_entrega = calcular_data_final_uteis(agora, prazo).strftime(
            "%d/%m/%Y"
        )

        st.markdown(
            "<h4 style='color: #003366;'>🔍 Seleção de Descontos (Limite Máximo de 15%)</h4>",
            unsafe_allow_html=True,
        )

        c1, c2, c3 = st.columns(3)
        desc_me = c1.checkbox("Microempreendedor com CNPJ (2.5%)")
        desc_nova = c1.checkbox("Empresa nova no mercado (2.5%)")
        desc_antigo = c1.checkbox("Cliente antigo da EJ (5.0%)")
        desc_novo_cli = c2.checkbox("Primeira compra / Novo cliente (1.25%)")
        desc_mulher = c2.checkbox("Liderança feminina / Mulheres (1.75%)")
        desc_combo = c2.checkbox("Combo / Mais de um serviço (5.0%)")
        desc_qtd = c3.checkbox("Quantidade progressiva de produtos (5.0%)")
        desc_indica = c3.checkbox("Lead por indicação de cliente (5.0%)")
        desc_ufmg = c3.checkbox("Ex-aluno ou familiar UFMG (2.5%)")

        soma_descontos = 0
        motivos = []
        if desc_me:
            soma_descontos += 2.5
            motivos.append("Microempreendedor")
        if desc_nova:
            soma_descontos += 2.5
            motivos.append("Empresa nova")
        if desc_antigo:
            soma_descontos += 5.0
            motivos.append("Cliente antigo da EJ")
        if desc_novo_cli:
            soma_descontos += 1.25
            motivos.append("Novos clientes")
        if desc_mulher:
            soma_descontos += 1.75
            motivos.append("Líderes mulheres")
        if desc_combo:
            soma_descontos += 5.0
            motivos.append("Mais de um serviço")
        if desc_qtd:
            soma_descontos += 5.0
            motivos.append("Quantidade de produto")
        if desc_indica:
            soma_descontos += 5.0
            motivos.append("Indicação de cliente")
        if desc_ufmg:
            soma_descontos += 2.5
            motivos.append("Vínculo UFMG")

        desconto_final_pct = min(soma_descontos, 15.0)
        motivo_txt = (
            ", ".join(motivos) if motivos else "Critérios de elegibilidade"
        )

        total_desconto = total_cheio * (1 - (desconto_final_pct / 100))
        unitario_desconto = total_desconto / quantidade

        parcelas_desconto = obter_texto_parcelamento(
            nome_servico, total_desconto
        )

        st.markdown(
            f"**Soma dos descontos:** {soma_descontos:.2f}% | **Desconto real"
            " aplicado (Limite de 15%):** <span style='color:#003366;"
            f" font-weight:bold;'>{desconto_final_pct:.2f}%</span>",
            unsafe_allow_html=True,
        )

        st.markdown("---")
        card1, card2 = st.columns(2)
        card1.metric(
            "Opção 1: Preço Cheio",
            f"R$ {total_cheio:,.2f}",
            f"{parcelas_cheio}",
        )
        card2.metric(
            "Opção 2: Preço com Desconto",
            f"R$ {total_desconto:,.2f}",
            f"{parcelas_desconto}",
        )

        dados_calculados = {
            "nome_lead": nome_lead,
            "tipo_servico": tipo_servico,
            "nome_servico": nome_servico,
            "valor_base": valor_base,
            "quantidade": quantidade,
            "total_cheio": total_cheio,
            "prazo": prazo,
            "data_entrega_autoral": data_entrega,
            "parcelas_cheio": parcelas_cheio,
            "parcelas_desconto": parcelas_desconto,
            "motivos_lista": motivos,
            "motivo_desconto": motivo_txt,
            "unitario_desconto": unitario_desconto,
            "total_desconto": total_desconto,
        }

    else:
        # --- CENÁRIO TERCEIRIZADO ---
        st.markdown(
            "<h4 style='color: #003366;'>📂 Upload do Orçamento do Laboratório</h4>",
            unsafe_allow_html=True,
        )
        arquivo_pdf = st.file_uploader(
            "Arraste o arquivo PDF do laboratório parceiro aqui:",
            type=["pdf"],
            key="file_pdf_lab",
        )

        val_auto_detectado = 0.0

        if arquivo_pdf is not None:
            with st.spinner(
                "🤖 IA processando o PDF e extraindo o valor cobrado..."
            ):
                valor_extraido = extrair_valor_pdf_com_ia(arquivo_pdf)
                if valor_extraido is not None and valor_extraido > 0:
                    val_auto_detectado = valor_extraido
                    st.success(
                        "✅ Processado com sucesso! Valor base do laboratório"
                        f" identificado: **R$ {val_auto_detectado:,.2f}**"
                    )

        orcamento_lab = st.number_input(
            "Valor do Orçamento do Laboratório (R$):",
            min_value=0.0,
            value=float(val_auto_detectado),
            step=50.0,
        )

        col1, col2 = st.columns(2)
        with col1:
            nome_servico = st.text_input(
                "Nome do Serviço Laboratorial:",
                value="Análise Microbiológica de Água",
            )
        with col2:
            prazo_lab = st.number_input(
                "Prazo original dado pelo laboratório (em dias corridos):",
                min_value=1,
                value=7,
            )
            coleta_opcao = st.radio(
                "Serviço de Coleta:",
                [
                    "Não oferecido (Amostra por conta do cliente)",
                    "Oferecido pela Farmácia Jr.",
                ],
            )
            valor_coleta = (
                st.number_input(
                    "Valor da Coleta (R$):", min_value=0.0, value=0.0
                )
                if "Oferecido" in coleta_opcao
                else 0.0
            )

        parametros = st.text_area(
            "Parâmetros analisados:",
            value=(
                "Presença/Ausência de bactérias do grupo Coliformes Totais e"
                " Termotolerantes (E. coli)."
            ),
        )
        metodologia = st.text_area(
            "Metodologia aplicada:",
            value=(
                "Contagem de bactérias heterotróficas conforme padrões"
                " laboratoriais."
            ),
        )

        # Matemática da planilha: = B4 + 110 + 0.18*(B4 + 110)
        margem_fixa_setor = 110.00
        subtotal_terc = orcamento_lab + margem_fixa_setor
        taxas_de_nota = subtotal_terc * 0.18
        total_terceirizado = subtotal_terc + taxas_de_nota

        prazo_final_terc = int(round(prazo_lab + 20, 0))
        data_entrega_terc = (agora + timedelta(days=prazo_final_terc)).strftime(
            "%d/%m/%Y"
        )

        parcelas_terceirizado = obter_texto_parcelamento(
            "Tabela Nutricional", total_terceirizado
        )

        txt_coleta = (
            f"R$ {valor_coleta:,.2f}"
            if "Oferecido" in coleta_opcao
            else (
                "não é oferecido o serviço de coleta ficando a critério do"
                " cliente a disponibilização da amostra."
            )
        )

        st.markdown("---")
        st.metric(
            "Opção Única Terceirizada (Cálculo Automático)",
            f"R$ {total_terceirizado:,.2f}",
            f"Prazo com segurança: {prazo_final_terc} dias corridos",
        )
        st.info(
            "💳 Condição de Pagamento Calculada:"
            f" **{parcelas_terceirizado}**"
        )
        st.info(
            "📅 Previsão exata de entrega para o cliente final:"
            f" **{data_entrega_terc}**"
        )

        dados_calculados = {
            "nome_lead": nome_lead,
            "tipo_servico": tipo_servico,
            "nome_servico": nome_servico,
            "total_terceirizado": total_terceirizado,
            "prazo_terceirizado": prazo_final_terc,
            "data_entrega_terc": data_entrega_terc,
            "parametros": parametros,
            "metodologia": metodologia,
            "txt_coleta": txt_coleta,
            "parcelas_terceirizado": parcelas_terceirizado,
        }

    if nome_lead:
        st.markdown("---")
        arquivo_word = gerar_docx_proposta(dados_calculados)
        st.download_button(
            label=(
                "📥 Gerar e Baixar Documento de Precificação Oficial (.docx)"
            ),
            data=arquivo_word,
            file_name=f"Precificação - {nome_lead}.docx",
            mime=(
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            ),
            use_container_width=True,
        )
    else:
        st.warning(
            "⚠️ Insira o nome do Lead no início da página para habilitar a"
            " geração do documento Word."
        )
